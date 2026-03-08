import os
import logging
from pathlib import Path
from typing import Optional
import pdfplumber
import pytesseract
from PIL import Image
import docx
import pypdf

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.jpg': self._extract_image,
            '.jpeg': self._extract_image,
            '.png': self._extract_image,
            '.txt': self._extract_text,
            '.docx': self._extract_docx
        }
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from a document based on its file type"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            logger.info(f"Extracting text from {file_path}")
            text = self.supported_formats[file_extension](file_path)
            logger.info(f"Extracted {len(text)} characters from {file_path}")
            return text
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using pdfplumber"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path}: {e}")
            # Fallback to pypdf
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e2:
                logger.error(f"Both PDF extractors failed for {file_path}: {e2}")
                raise
        
        return text.strip()
    
    def _extract_image(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            raise
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Failed to read text file {file_path}: {e}")
                raise
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract from DOCX {file_path}: {e}")
            raise
    
    def get_document_info(self, file_path: str) -> dict:
        """Get basic information about the document"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        stat = file_path.stat()
        
        return {
            "filename": file_path.name,
            "size": stat.st_size,
            "extension": file_path.suffix.lower(),
            "modified": stat.st_mtime,
            "is_supported": file_path.suffix.lower() in self.supported_formats
        }
